"""
Подготовка шаблонов из оригинальных DOCX.
ВАЖНО: Не трогает раны с floating-изображениями (anchor).
"""
import shutil
import os
from docx import Document

WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'


def _has_drawing(run) -> bool:
    """Ран содержит floating или inline изображение."""
    el = run._element
    return (
        bool(el.findall('.//{%s}anchor' % WP_NS)) or
        bool(el.findall('.//{%s}inline' % WP_NS))
    )


def _merge_text_runs(para):
    """
    Объединяет только текстовые раны (без картинок) в первый текстовый ран.
    Раны с картинками НЕ трогает — они остаются на своих местах.
    """
    # Найти индекс первого текстового рана (без drawing)
    first_text_idx = None
    for i, run in enumerate(para.runs):
        if not _has_drawing(run):
            first_text_idx = i
            break

    if first_text_idx is None:
        return  # только картинки — не трогаем

    # Собрать весь текст из текстовых ранов
    full_text = ''.join(
        r.text for r in para.runs
        if not _has_drawing(r) and r.text
    )

    # Нормализовать неразрывные пробелы
    full_text = full_text.replace('\u00a0', ' ').replace('\xa0', ' ')

    # Записать весь текст в первый текстовый ран
    para.runs[first_text_idx].text = full_text

    # Обнулить все остальные текстовые раны (картинки не трогаем)
    for i, run in enumerate(para.runs):
        if i != first_text_idx and not _has_drawing(run):
            run.text = ''


def _all_paras(doc):
    """Все параграфы: body + таблицы."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def merge_all(doc):
    for p in _all_paras(doc):
        _merge_text_runs(p)


def replace_all(doc, old: str, new: str):
    """Заменить все вхождения old→new во всех текстовых ранах."""
    for p in _all_paras(doc):
        for run in p.runs:
            if _has_drawing(run):
                continue
            if run.text and old in run.text:
                run.text = run.text.replace(old, new)


def replace_first(doc, old: str, new: str) -> bool:
    """Заменить первое вхождение old→new. Возвращает True если нашёл."""
    for p in _all_paras(doc):
        for run in p.runs:
            if _has_drawing(run):
                continue
            if run.text and old in run.text:
                run.text = run.text.replace(old, new, 1)
                return True
    return False


def replace_ordered(doc, replacements: list):
    for old, new in replacements:
        replace_all(doc, old, new)


# ─── Замены ───────────────────────────────────────────────────────────────────

BASE_REPLACEMENTS = [
    # ── Сделка ────────────────────────────────────────────────────────────────
    ("VO 99082 За виртуальный актив по заявке №499 от 17.02.2026",
     "VO {{KVVO}} За виртуальный актив по заявке №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("Покупка виртуальных активов по заявке №499 от 17.02.2026",
     "{{OPERATION_TYPE}} по заявке №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("Продажа виртуальных активов по заявке №499 от 17.02.2026",
     "{{OPERATION_TYPE}} по заявке №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("Договор №499 от 17.02.2026",   "Договор №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("Акт №499 от 17.02.2026",       "Акт №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("Заявка №499 от 17.02.2026",    "Заявка №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("по заявке №499 от 17.02.2026", "по заявке №{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("№499  от 17.02.2026",          "№{{DEAL_NUMBER}} от {{DEAL_DATE}}"),  # двойной пробел
    ("№499 от 17.02.2026",           "№{{DEAL_NUMBER}} от {{DEAL_DATE}}"),
    ("«17» февраля 2026 г.",         "{{DEAL_DATE_FULL}}"),
    ("17.02.2026",                   "{{DEAL_DATE}}"),

    # ── Справка: полные предложения (до коротких замен ИНН!) ──────────────────
    ('Общество с ограниченной ответственностью "СТЕЙБЛЕКС", ИНН 9909730748, Местонахождение: Кыргызская Республика, город Бишкек, Свердловский район, улица Московская, дом 197',
     '{{OP_FULL_NAME}}, ИНН {{OP_INN}}, Местонахождение: {{OP_ADDRESS_FULL}}'),
    ('Общество с ограниченной ответственностью "Алтынкопрю", ИНН 9909745705, Местонахождение: Кыргызстан, г. Бишкек, ул. Целинная 47',
     '{{CL_FULL_NAME}}, ИНН {{CL_INN}}, Местонахождение: {{CL_ADDRESS}}'),

    # ── Оператор ──────────────────────────────────────────────────────────────
    ("Общество с ограниченной ответственностью «СТЕЙБЛЕКС»", "{{OP_FULL_NAME}}"),
    ('Общество с ограниченной ответственностью "СТЕЙБЛЕКС"', "{{OP_FULL_NAME}}"),
    ("ОсОО «СТЕЙБЛЕКС»",    "{{OP_SHORT_NAME}}"),
    ("ОсОО «Стейблекс»",    "{{OP_SHORT_NAME}}"),
    ("ОсОО «Стеблекс»",     "{{OP_SHORT_NAME}}"),
    ("ИНН:9909730748",       "ИНН:{{OP_INN}}"),
    ("ИНН: 9909730748",      "ИНН: {{OP_INN}}"),
    ("ИНН 9909730748",       "ИНН {{OP_INN}}"),
    ("КПП:770387001",        "КПП:{{OP_KPP}}"),
    ("КПП: 770387001",       "КПП: {{OP_KPP}}"),
    ("693010, г.Южно-Сахалинск, ул Комсомольская, 145", "{{OP_LEGAL_ADDRESS}}"),
    ("720009, г. Бишкек, ул. Московская, д. 197",       "{{OP_ADDRESS}}"),
    ("Кыргызская Республика, город Бишкек, Свердловский район, улица Московская, дом 197",
     "{{OP_ADDRESS_FULL}}"),
    ("ИНН: 02504202410133",  "ИНН: {{OP_INN}}"),
    ("02504202410133",       "{{OP_INN}}"),
    ("Лицензия: 150 от 28-03-2025", "Лицензия: {{OP_LICENSE}}"),
    ("Зенков И.В.",          "{{OP_DIRECTOR}}"),

    # ── Клиент ────────────────────────────────────────────────────────────────
    ("Общество с ограниченной ответственностью «Алтынкопрю»", "{{CL_FULL_NAME}}"),
    ('Общество с ограниченной ответственностью "Алтынкопрю"', "{{CL_FULL_NAME}}"),
    ("ОсОО «Алтынкопрю»",   "{{CL_SHORT_NAME}}"),
    ("ИНН: 9909745705",      "ИНН: {{CL_INN}}"),
    ("ИНН:9909745705",       "ИНН:{{CL_INN}}"),
    ("ИНН 9909745705",       "ИНН {{CL_INN}}"),
    ("КПП: 770887001",       "КПП: {{CL_KPP}}"),
    ("КПП:770887001",        "КПП:{{CL_KPP}}"),
    ("317744-3301-ООО",      "{{CL_REG_NUMBER}}"),
    ("Кыргызстан, г. Бишкек, ул. Целинная 47", "{{CL_ADDRESS}}"),

    # ── Счета (до банка!) ─────────────────────────────────────────────────────
    ("40807810500014264602", "{{OP_BANK_ACCOUNT}}"),
    ("40807810600014672000", "{{CL_BANK_ACCOUNT}}"),

    # ── Кошельки ──────────────────────────────────────────────────────────────
    ("TXFEYN4C5BnesaxUXJiXJHGS7K12QutZ3r", "{{OP_WALLET}}"),
    ("THtSiaKaPF1R1dhZpBAcmgkvchnDmoA9Pi", "{{CL_WALLET}}"),

    # ── КВВО ──────────────────────────────────────────────────────────────────
    ("99082", "{{KVVO}}"),
    ("73074", "{{CL_KIO_PLACEHOLDER}}"),

    # ── Суммы ─────────────────────────────────────────────────────────────────
    ("68 883.559500 RUB",  "{{FIAT_AMOUNT_SHORT}} RUB"),
    ("500 USDT",           "{{VA_AMOUNT_SHORT}} {{VA_TYPE}}"),
    ("6 287.726 USDT",     "{{VA_AMOUNT}} {{VA_TYPE}}"),
    ("6287.726",           "{{VA_AMOUNT}}"),
    ("500 000  RUR",       "{{FIAT_AMOUNT}} RUR"),   # двойной пробел
    ("500 000 RUR",        "{{FIAT_AMOUNT}} RUR"),
    ("500 000 RUB",        "{{FIAT_AMOUNT}} RUB"),
    ("500 000",            "{{FIAT_AMOUNT}}"),
    ("500000",             "{{FIAT_AMOUNT}}"),
    ("79,52",              "{{EXCHANGE_RATE}}"),
    ("USDT_TRC20",         "{{VA_TICKER}}"),
    ("TRC-20",             "{{NETWORK}}"),
    ("d30791fb7a5a460fce1ff756e0467aff26efea899c265d51abcb247af50f31e6", "{{TX_HASH}}"),
]

SELL_EXTRA = [("{{CL_KIO_PLACEHOLDER}}", "-")]
BUY_EXTRA  = [("{{CL_KIO_PLACEHOLDER}}", "{{OP_KIO}}")]


def apply_bank_placeholders(doc):
    """
    Банк и БИК одинаковы у обоих участников.
    Первое вхождение → клиент, второе → оператор.
    """
    bank_str = 'КБ "Долинск" (АО)'
    replace_first(doc, bank_str, "{{CL_BANK_NAME}}")
    replace_first(doc, bank_str, "{{OP_BANK_NAME}}")
    replace_first(doc, "046401727", "{{CL_BANK_BIK}}")
    replace_first(doc, "046401727", "{{OP_BANK_BIK}}")


def verify(doc, template_name: str):
    """Проверяет что ключевые плейсхолдеры созданы и нет старых значений."""
    all_text = ''
    for p in _all_paras(doc):
        for r in p.runs:
            if not _has_drawing(r) and r.text:
                all_text += r.text

    must_have = [
        '{{DEAL_NUMBER}}', '{{FIAT_AMOUNT}}', '{{VA_AMOUNT}}',
        '{{CL_BANK_NAME}}', '{{OP_BANK_NAME}}', '{{CL_INN}}', '{{OP_INN}}',
    ]
    must_not  = ['499', '9909745705', '9909730748', '500 000', '6287']

    ok = True
    for ph in must_have:
        if ph not in all_text:
            print(f'  ⚠ НЕ СОЗДАН: {ph}')
            ok = False
    for val in must_not:
        if val in all_text:
            print(f'  ⚠ ОСТАЛСЯ ОРИГИНАЛ: {val}')
            ok = False

    # Проверить что картинки на месте
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    drawing_count = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if _has_drawing(r):
                drawing_count += 1
    print(f'  Floating-картинок в шаблоне: {drawing_count}')

    if ok:
        print(f'  ✅ {template_name}: все плейсхолдеры OK')
    return ok


def prepare(src: str, dst: str, extra: list):
    print(f"  {os.path.basename(src)} → {os.path.basename(dst)}")
    shutil.copy(src, dst)
    doc = Document(dst)

    # Сколько картинок в оригинале
    wp_ns = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    orig_drawings = sum(
        1 for p in doc.paragraphs
        for r in p.runs
        if r._element.findall('.//{%s}anchor' % wp_ns)
    )
    print(f"  Картинок в оригинале: {orig_drawings}")

    merge_all(doc)
    replace_ordered(doc, BASE_REPLACEMENTS + extra)
    apply_bank_placeholders(doc)
    doc.save(dst)

    # Reload and verify
    doc2 = Document(dst)
    verify(doc2, os.path.basename(dst))
    print(f"  ✓ сохранён: {dst}")


if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    tpl  = os.path.join(base, "templates")

    print("Подготовка шаблона SELL...")
    prepare(
        src=os.path.join(tpl, "Акт_покупке_у_нас_ЮЛ.docx"),
        dst=os.path.join(tpl, "template_sell.docx"),
        extra=SELL_EXTRA,
    )

    print("\nПодготовка шаблона BUY...")
    prepare(
        src=os.path.join(tpl, "Акт_продаже_нам_от_ЮЛ.docx"),
        dst=os.path.join(tpl, "template_buy.docx"),
        extra=BUY_EXTRA,
    )

    print("\n✅ Шаблоны готовы")
