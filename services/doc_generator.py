"""
Генерация DOCX из шаблона.
Заменяет плейсхолдеры по одному рану — не трогает раны с картинками.
"""
import os
import copy
import shutil
from datetime import datetime
from docx import Document
from config import TEMPLATES_DIR, OUTPUT_DIR, ACT_TEMPLATE_FILES
from utils import safe_name_part

WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

MONTHS_RU = {
    1:"января",2:"февраля",3:"марта",4:"апреля",
    5:"мая",6:"июня",7:"июля",8:"августа",
    9:"сентября",10:"октября",11:"ноября",12:"декабря",
}


def _has_drawing(run) -> bool:
    el = run._element
    return (
        bool(el.findall('.//{%s}anchor' % WP_NS)) or
        bool(el.findall('.//{%s}inline' % WP_NS))
    )


def _s(v) -> str:
    return str(v) if v is not None else ''


def _replace_in_doc(doc: Document, replacements: dict):
    """
    Заменяет плейсхолдеры в параграфе с учётом того, что в Word плейсхолдер
    может быть разбит на несколько ранов (например `{{OP_`, `KIO`, `}}`).
    Параграфы с картинками не склеиваются — там замена идёт поран­но.
    """
    def _flush(group):
        # Склеивает группу соседних текстовых ранов и заменяет в ней плейсхолдеры
        if not group:
            return
        gfull = ''.join(r.text or '' for r in group)
        if '{{' not in gfull:
            return
        new = gfull
        for k, v in replacements.items():
            if k in new:
                new = new.replace(k, _s(v))
        if new != gfull:
            group[0].text = new
            for r in group[1:]:
                r.text = ''

    def process_para(para):
        runs = para.runs
        if not runs:
            return
        if '{{' not in ''.join(r.text or '' for r in runs):
            return
        # Группируем подряд идущие текстовые раны; раны-картинки разделяют группы,
        # чтобы текст не «перепрыгивал» через изображение (подпись/печать).
        group = []
        for r in runs:
            if _has_drawing(r):
                _flush(group)
                group = []
            else:
                group.append(r)
        _flush(group)

    for para in doc.paragraphs:
        process_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)


def format_date_full(date_str: str) -> str:
    try:
        d = datetime.strptime(date_str, "%d.%m.%Y")
        return f"«{d.day}» {MONTHS_RU[d.month]} {d.year} г."
    except Exception:
        return date_str


def format_number(value: str) -> str:
    try:
        n = float(str(value).replace(',', '.').replace(' ', ''))
        if n == int(n):
            return f"{int(n):,}".replace(',', ' ')
        s = f"{n:.6f}".rstrip('0').rstrip('.')
        parts = s.split('.')
        parts[0] = f"{int(parts[0]):,}".replace(',', ' ')
        return '.'.join(parts)
    except Exception:
        return str(value)


def build_replacements(deal: dict, op: dict, client: dict,
                       client_custom: dict | None = None) -> dict:
    deal_date      = deal.get('deal_date', '')
    deal_date_full = format_date_full(deal_date)
    exec_date      = deal.get('execution_date', deal_date)
    va_amount_fmt  = format_number(deal.get('va_amount', '0'))
    fiat_amount_fmt= format_number(deal.get('fiat_amount', '0'))
    va_type        = deal.get('va_type', 'USDT')
    network        = deal.get('network', 'TRC-20')
    va_ticker      = f"{va_type}_{network.replace('-','')}"
    kvvo           = deal.get('kvvo', '99082')
    act_type       = deal.get('act_type', 'sell')
    deal_number    = str(deal.get('deal_number', ''))
    # 'buy' = клиент продаёт нам; 'sell'/'invoice' = клиент покупает у нас
    operation_type = (
        'Продажа виртуальных активов' if act_type == 'buy'
        else 'Покупка виртуальных активов'
    )

    result = {
        # Сделка
        '{{DEAL_NUMBER}}':       deal_number,
        '{{DEAL_DATE}}':         deal_date,
        '{{DEAL_DATE_FULL}}':    deal_date_full,
        '{{OPERATION_TYPE}}':    operation_type,
        '{{KVVO}}':              kvvo,
        '{{VA_TYPE}}':           va_type,
        '{{NETWORK}}':           network,
        '{{VA_TICKER}}':         va_ticker,
        '{{VA_AMOUNT}}':         va_amount_fmt,
        '{{VA_AMOUNT_SHORT}}':   va_amount_fmt,
        '{{FIAT_AMOUNT}}':       fiat_amount_fmt,
        '{{FIAT_AMOUNT_SHORT}}': fiat_amount_fmt,
        '{{FIAT_CURRENCY}}':     'RUB',
        '{{EXCHANGE_RATE}}':     str(deal.get('exchange_rate', '')).replace('.', ','),
        '{{COMMISSION_FIAT}}':   str(deal.get('commission_fiat', '0%')),
        '{{COMMISSION_VA}}':     str(deal.get('commission_va', '0%')),
        '{{TX_HASH}}':           deal.get('tx_hash', ''),
        '{{EXECUTION_DATE}}':    exec_date,
        '{{CL_WALLET}}':         deal.get('client_wallet', client.get('wallet', '')),
        '{{OP_WALLET}}':         deal.get('operator_wallet', op.get('wallet', '')),
        # Оператор
        '{{OP_FULL_NAME}}':      op.get('full_name', ''),
        '{{OP_SHORT_NAME}}':     op.get('short_name', ''),
        '{{OP_INN}}':            op.get('inn', ''),
        '{{OP_KPP}}':            op.get('kpp', ''),
        '{{OP_ADDRESS}}':        op.get('address', ''),
        '{{OP_ADDRESS_FULL}}':   op.get('address', ''),
        '{{OP_LEGAL_ADDRESS}}':  op.get('legal_address', op.get('address', '')),
        '{{OP_LICENSE}}':        op.get('license', ''),
        '{{OP_KIO}}':            op.get('kio', ''),
        '{{OP_INN_RF}}':         op.get('inn_rf', op.get('inn', '')),
        '{{OP_KPP_RF}}':         op.get('kpp_rf', op.get('kpp', '')),
        '{{OP_BANK_NAME}}':      op.get('bank_name', ''),
        '{{OP_BANK_ACCOUNT}}':   op.get('bank_account', ''),
        '{{OP_BANK_BIK}}':       op.get('bank_bik', ''),
        '{{OP_DIRECTOR}}':       op.get('director_name', ''),
        # Клиент
        '{{CL_FULL_NAME}}':      client.get('full_name', ''),
        '{{CL_SHORT_NAME}}':     client.get('short_name', ''),
        '{{CL_INN}}':            client.get('inn', ''),
        '{{CL_KPP}}':            client.get('kpp', ''),
        '{{CL_REG_NUMBER}}':     client.get('reg_number', ''),
        '{{CL_ADDRESS}}':        client.get('address', ''),
        '{{CL_KIO}}':            client.get('kio', '-'),
        '{{CL_INN_RF}}':         client.get('inn_rf', '-'),
        '{{CL_KPP_RF}}':         client.get('kpp_rf', '-'),
        '{{CL_BANK_NAME}}':      client.get('bank_name', ''),
        '{{CL_BANK_ACCOUNT}}':   client.get('bank_account', ''),
        '{{CL_BANK_BIK}}':       client.get('bank_bik', ''),
    }
    # Свои переменные: любой ключ реквизитов оператора доступен как {{OP_KEY}},
    # значения пользовательских полей компании — как {{CL_KEY}}.
    for k, v in (op or {}).items():
        result.setdefault('{{OP_%s}}' % k.upper(), _s(v))
    for k, v in (client_custom or {}).items():
        result.setdefault('{{CL_%s}}' % k.upper(), _s(v))
    return result


def _to_float(value) -> float:
    try:
        return float(str(value).replace(',', '.').replace(' ', ''))
    except Exception:
        return 0.0


def get_transactions(deal: dict) -> list:
    """
    Список транзакций (частей сделки). Несколько — только если их явно
    добавили (>1). Для одной берём текущие поля сделки (так работает ручное
    редактирование суммы/хэша через меню правок).
    """
    txs = deal.get('transactions')
    if txs and len(txs) > 1:
        return txs
    return [{
        'va_amount': deal.get('va_amount', '0'),
        'tx_hash':   deal.get('tx_hash', ''),
    }]


def computed_transactions(deal: dict) -> list:
    """
    Транзакции-части сделки с рассчитанными рублями по курсу.
    Сумма ВА частей = общей сумме сделки; рубли каждой части = ВА × курс,
    последняя часть (остаток) добивается так, чтобы рубли точно дали общую сумму.
    Возвращает [{'va_amount': str, 'fiat_amount': float, 'tx_hash': str}, ...].
    """
    txs        = get_transactions(deal)
    total_va   = sum(_to_float(t.get('va_amount', 0)) for t in txs)
    total_fiat = _to_float(deal.get('fiat_amount', 0))
    rate       = (total_fiat / total_va) if total_va else 0.0

    out, acc = [], 0.0
    last = len(txs) - 1
    for i, t in enumerate(txs):
        va = _to_float(t.get('va_amount', 0))
        if i < last:
            fiat = round(va * rate, 2)
            acc += fiat
        else:
            fiat = round(total_fiat - acc, 2)   # остаток добивает до точной общей суммы
        out.append({
            'va_amount':   t.get('va_amount', '0'),
            'fiat_amount': fiat,
            'tx_hash':     t.get('tx_hash', ''),
        })
    return out


def _row_text(row) -> str:
    return ''.join(cell.text for cell in row.cells)


def _merge_row_runs(row):
    """Склеивает раны в каждой ячейке строки — плейсхолдеры становятся целыми."""
    for cell in row.cells:
        for para in cell.paragraphs:
            runs = [r for r in para.runs if not _has_drawing(r)]
            if not runs:
                continue
            full = ''.join(r.text or '' for r in runs)
            if full:
                runs[0].text = full
                for r in runs[1:]:
                    r.text = ''


def _fill_tx_row(row, tx: dict):
    """Подставляет в строку-транзакцию её сумму ВА, рубли (по курсу) и хэш."""
    va   = format_number(tx.get('va_amount', '0'))
    fiat = format_number(tx.get('fiat_amount', '0'))
    mapping = {
        '{{VA_AMOUNT}}':         va,
        '{{VA_AMOUNT_SHORT}}':   va,
        '{{FIAT_AMOUNT}}':       fiat,
        '{{FIAT_AMOUNT_SHORT}}': fiat,
        '{{TX_HASH}}':           tx.get('tx_hash', '') or '',
    }
    for cell in row.cells:
        for para in cell.paragraphs:
            for r in para.runs:
                if _has_drawing(r) or not r.text:
                    continue
                for k, v in mapping.items():
                    if k in r.text:
                        r.text = r.text.replace(k, v)


def _expand_transactions(doc: Document, transactions: list):
    """
    Находит таблицу транзакций (строку с {{TX_HASH}}) и клонирует строку данных
    по числу транзакций. Строка «Итого» остаётся для подстановки суммарных значений.
    """
    for table in doc.tables:
        data_idx = None
        for i, row in enumerate(table.rows):
            if '{{TX_HASH}}' in _row_text(row):
                data_idx = i
                break
        if data_idx is None:
            continue

        data_row = table.rows[data_idx]
        _merge_row_runs(data_row)

        # Клонируем строку данных для второй и последующих транзакций
        tr = data_row._tr
        last = tr
        for _ in range(len(transactions) - 1):
            new_tr = copy.deepcopy(tr)
            last.addnext(new_tr)
            last = new_tr

        for j, tx in enumerate(transactions):
            _fill_tx_row(table.rows[data_idx + j], tx)

        # Строка «Итого» идёт сразу после блока транзакций
        total_idx = data_idx + len(transactions)
        if total_idx < len(table.rows):
            _merge_row_runs(table.rows[total_idx])
        return  # таблица транзакций одна


def scan_placeholders(docx_path: str) -> list:
    """
    Возвращает все плейсхолдеры {{...}} из шаблона (склеивая раны,
    чтобы найти и разбитые на части).
    """
    import re
    doc = Document(docx_path)
    texts = []

    def collect(paragraphs):
        for para in paragraphs:
            texts.append(''.join(r.text or '' for r in para.runs))

    collect(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                collect(cell.paragraphs)
    for section in doc.sections:
        for part in (section.header, section.footer):
            if part is not None:
                collect(part.paragraphs)

    found = re.findall(r'\{\{[A-Za-z0-9_]+\}\}', '\n'.join(texts))
    seen, out = set(), []
    for p in found:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out


def known_placeholders(op: dict, custom_fields: list) -> set:
    """Множество плейсхолдеров, которые бот умеет заполнять."""
    base = set(build_replacements({}, op or {}, {}).keys())
    base |= {'{{CL_%s}}' % f['key'].upper() for f in (custom_fields or [])}
    return base


def generate_docx(deal: dict, op: dict, client: dict,
                  client_custom: dict | None = None) -> str:
    act_type = deal.get('act_type', 'sell')
    # Соответствие типа операции файлу шаблона — в config.ACT_TEMPLATE_FILES
    template_name = ACT_TEMPLATE_FILES.get(act_type, ACT_TEMPLATE_FILES['sell'])
    direction = {'sell': 'buy', 'buy': 'sell', 'invoice': 'invoice_buy'}.get(act_type, act_type)
    template_path = os.path.join(TEMPLATES_DIR, template_name)

    if not os.path.exists(template_path):
        raise FileNotFoundError(
            f"Шаблон не найден: {template_path}\n"
            f"Положите файл {template_name} в папку {TEMPLATES_DIR}/"
        )

    # Транзакции-части с рассчитанными рублями; общая сумма ВА = сумме частей
    transactions = computed_transactions(deal)
    va_total = sum(_to_float(t.get('va_amount', 0)) for t in transactions)

    deal_totals = dict(deal)
    deal_totals['va_amount'] = va_total   # «Итого» по ВА = сумма всех частей
    deal_totals['tx_hash']   = ''         # хэши подставляются построчно
    # fiat_amount остаётся введённой общей суммой → идёт в «Итого» по рублям

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    deal_number = deal.get('deal_number', '0')
    deal_date   = deal.get('deal_date', '').replace('.', '')
    cl_short    = safe_name_part(client.get('short_name', ''), limit=10, default='CL')
    out_name    = f"act_{direction}_{deal_number}-{cl_short}_{deal_date}.docx"
    out_path    = os.path.join(OUTPUT_DIR, out_name)

    shutil.copy(template_path, out_path)
    doc          = Document(out_path)
    # 1) построчно заполняем таблицу транзакций (до общей замены)
    _expand_transactions(doc, transactions)
    # 2) общая замена всех остальных плейсхолдеров (сводные суммы, реквизиты)
    replacements = build_replacements(deal_totals, op, client, client_custom)
    _replace_in_doc(doc, replacements)
    doc.save(out_path)
    return out_path
